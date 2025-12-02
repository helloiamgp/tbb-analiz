#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
TBB Yazı Otomatik Yanıtlama Sistemi
Aytemiz Yatırım Bankası A.Ş.
Versiyon: 1.0 Final
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import os
import re
import sys
from datetime import datetime

# Hata yakalama ile import
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

# PDF kütüphaneleri
PDF_LIB = None
try:
    import fitz
    PDF_LIB = 'fitz'
except ImportError:
    try:
        import pdfplumber
        PDF_LIB = 'pdfplumber'
    except ImportError:
        try:
            from PyPDF2 import PdfReader
            PDF_LIB = 'PyPDF2'
        except ImportError:
            pass


class TBBYanitSistemi:
    def __init__(self, root):
        self.root = root
        self.root.title("TBB Yazı Otomatik Yanıtlama Sistemi - Aytemiz Yatırım Bankası")
        self.root.geometry("1200x900")
        self.root.configure(bg='#f0f0f0')
        
        # Icon ayarla (varsa)
        try:
            self.root.iconbitmap('icon.ico')
        except:
            pass
        
        # Başlangıç kontrolleri
        if not DOCX_AVAILABLE:
            messagebox.showerror("Hata", "python-docx kütüphanesi bulunamadı!\nLütfen: pip install python-docx")
            sys.exit(1)
        
        if not PANDAS_AVAILABLE:
            messagebox.showerror("Hata", "pandas kütüphanesi bulunamadı!\nLütfen: pip install pandas openpyxl")
            sys.exit(1)
        
        # Değişkenler
        self.musteri_listesi = None
        self.toplu_dosyalar = []
        self.toplu_sonuclar = []
        self.toplu_yanitlar = []
        
        # Arayüz oluştur
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        self.tekli_frame = tk.Frame(self.notebook, bg='#f0f0f0')
        self.notebook.add(self.tekli_frame, text="📄 Tekli İşlem")
        
        self.toplu_frame = tk.Frame(self.notebook, bg='#f0f0f0')
        self.notebook.add(self.toplu_frame, text="📁 Toplu İşlem")
        
        self.tekli_arayuz_olustur()
        self.toplu_arayuz_olustur()
        
        # Durum çubuğu
        durum_frame = tk.Frame(self.root, bg='#e0e0e0')
        durum_frame.pack(fill=tk.X, side=tk.BOTTOM)
        
        self.durum_label = tk.Label(durum_frame, text="Hazır", font=('Segoe UI', 9),
                                   bg='#e0e0e0', anchor=tk.W, padx=10)
        self.durum_label.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        pdf_durum = "✅ PDF desteği aktif" if PDF_LIB else "⚠️ PDF desteği yok"
        tk.Label(durum_frame, text=pdf_durum, font=('Segoe UI', 8), 
                bg='#e0e0e0', fg='#666', padx=10).pack(side=tk.RIGHT)
        
        # Otomatik müşteri listesi yükleme
        self.otomatik_musteri_yukle()

    # ==================== BİLGİ ÇIKARMA ====================
    
    def bilgi_cikar(self, icerik):
        """Yazıdan tüm gerekli bilgileri çıkar"""
        bilgiler = {
            'muhatap_kurum': '', 'muhatap_alt1': '', 'muhatap_alt2': '',
            'tarih': '', 'sayi': '', 'tckn': '', 'vkn': '', 'adsoyad': ''
        }
        
        satirlar = icerik.split('\n')
        
        # Muhatap kurumları bul
        kurum_anahtar = ['başkanlığı', 'müdürlüğü', 'dairesi', 'komutanlığı', 
                        'savcılığı', 'mahkemesi', 'kaymakamlığı', 'valiliği',
                        'defterdarlığı', 'bakanlığı', 'kurumu', 'idaresi']
        
        kurum_satirlari = []
        for satir in satirlar[:15]:
            satir_temiz = satir.strip()
            if satir_temiz and len(satir_temiz) > 3:
                if any(k in satir_temiz.lower() for k in kurum_anahtar):
                    kurum_satirlari.append(satir_temiz)
        
        if len(kurum_satirlari) >= 1:
            bilgiler['muhatap_kurum'] = kurum_satirlari[0]
        if len(kurum_satirlari) >= 2:
            bilgiler['muhatap_alt1'] = kurum_satirlari[1]
        if len(kurum_satirlari) >= 3:
            bilgiler['muhatap_alt2'] = kurum_satirlari[2]
        
        # Tarih bul
        tarih_match = re.search(r'(\d{2}[./]\d{2}[./]\d{4})', icerik)
        if tarih_match:
            bilgiler['tarih'] = tarih_match.group(1).replace('/', '.')
        
        # Sayı numarası bul
        sayi_patterns = [
            r'[Ss]ayı\s*:\s*([A-Za-z0-9\-\[\]\(\)\s\.\/]+?)(?:\n|$)',
            r'E-\d+[-\.\d\[\]]+',
            r'\d{5,}[-\.\d\[\]]+\s*[-–]\s*\d+',
        ]
        for pattern in sayi_patterns:
            sayi_match = re.search(pattern, icerik)
            if sayi_match:
                sayi = sayi_match.group(0) if pattern.startswith('E-') or pattern.startswith('\\d') else sayi_match.group(1)
                if len(sayi.strip()) > 5:
                    bilgiler['sayi'] = sayi.strip()
                    break
        
        # TCKN bul (11 haneli)
        tckn_match = re.search(r'\b(\d{11})\b', icerik)
        if tckn_match:
            bilgiler['tckn'] = tckn_match.group(1)
        
        # VKN bul (10 haneli)
        vkn_matches = re.findall(r'\b(\d{10})\b', icerik)
        for match in vkn_matches:
            if match != bilgiler['tckn'][:10] if bilgiler['tckn'] else True:
                bilgiler['vkn'] = match
                break
        
        # Ad Soyad bul
        if bilgiler['tckn']:
            pattern = rf"{bilgiler['tckn']}\s*(?:T\.?C\.?\s*)?(?:Kimlik\s*)?(?:Numaralı)?\s*([A-ZÇĞİÖŞÜa-zçğıöşü]+\s+[A-ZÇĞİÖŞÜa-zçğıöşü]+)"
            match = re.search(pattern, icerik)
            if match:
                bilgiler['adsoyad'] = match.group(1).strip()
        
        if not bilgiler['adsoyad']:
            match = re.search(r'([A-ZÇĞİÖŞÜ][a-zçğıöşü]+\s+[A-ZÇĞİÖŞÜ][A-Za-zçğıöşü]+)(?:\s+adlı|\s+isimli|\'[ıiuü]n)', icerik)
            if match:
                bilgiler['adsoyad'] = match.group(1).strip()
        
        return bilgiler

    # ==================== DOSYA OKUMA ====================
    
    def pdf_oku(self, dosya_yolu):
        """PDF dosyasını oku"""
        if not PDF_LIB:
            return "[PDF desteği yok. PDF okumak için PyMuPDF gerekli.]"
        
        metin = ""
        try:
            if PDF_LIB == 'fitz':
                import fitz
                doc = fitz.open(dosya_yolu)
                for sayfa in doc:
                    metin += sayfa.get_text()
                doc.close()
            elif PDF_LIB == 'pdfplumber':
                import pdfplumber
                with pdfplumber.open(dosya_yolu) as pdf:
                    for sayfa in pdf.pages:
                        t = sayfa.extract_text()
                        if t:
                            metin += t + "\n"
            elif PDF_LIB == 'PyPDF2':
                from PyPDF2 import PdfReader
                reader = PdfReader(dosya_yolu)
                for sayfa in reader.pages:
                    t = sayfa.extract_text()
                    if t:
                        metin += t + "\n"
        except Exception as e:
            return f"[PDF okuma hatası: {str(e)}]"
        
        return metin if metin.strip() else "[PDF'den metin çıkarılamadı]"
    
    def dosya_oku(self, dosya_yolu):
        """Dosya içeriğini oku"""
        uzanti = os.path.splitext(dosya_yolu)[1].lower()
        try:
            if uzanti == '.docx':
                doc = Document(dosya_yolu)
                return '\n'.join([p.text for p in doc.paragraphs])
            elif uzanti == '.pdf':
                return self.pdf_oku(dosya_yolu)
            elif uzanti == '.txt':
                with open(dosya_yolu, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            else:
                return f"[Desteklenmeyen format: {uzanti}]"
        except Exception as e:
            return f"[Dosya okuma hatası: {str(e)}]"

    # ==================== MÜŞTERİ LİSTESİ ====================
    
    def otomatik_musteri_yukle(self):
        """Aynı dizindeki müşteri listesini otomatik yükle"""
        # EXE'nin bulunduğu dizini bul
        if getattr(sys, 'frozen', False):
            uygulama_dizini = os.path.dirname(sys.executable)
        else:
            uygulama_dizini = os.path.dirname(os.path.abspath(__file__))
        
        for dosya in ["musteri_listesi.xlsx", "musteri_listesi.csv", "musteriler.xlsx"]:
            tam_yol = os.path.join(uygulama_dizini, dosya)
            if os.path.exists(tam_yol):
                self.musteri_listesi_yukle(tam_yol)
                break
    
    def musteri_listesi_sec(self):
        dosya = filedialog.askopenfilename(
            title="Müşteri Listesi Seçin",
            filetypes=[("Excel/CSV", "*.xlsx *.xls *.csv"), ("Tüm Dosyalar", "*.*")]
        )
        if dosya:
            self.musteri_listesi_yukle(dosya)
    
    def musteri_listesi_yukle(self, dosya_yolu):
        try:
            uzanti = os.path.splitext(dosya_yolu)[1].lower()
            if uzanti == '.csv':
                df = pd.read_csv(dosya_yolu, dtype=str)
            else:
                df = pd.read_excel(dosya_yolu, dtype=str)
            
            df.columns = df.columns.str.lower().str.strip()
            
            tckn_sutun = vkn_sutun = adsoyad_sutun = None
            for col in df.columns:
                c = col.lower()
                if 'tckn' in c or 'tc' in c or 'kimlik' in c:
                    tckn_sutun = col
                elif 'vkn' in c or 'vergi' in c:
                    vkn_sutun = col
                elif 'ad' in c or 'isim' in c or 'müşteri' in c or 'soyad' in c:
                    adsoyad_sutun = col
            
            if not tckn_sutun and not vkn_sutun:
                tckn_sutun = df.columns[0]
            
            self.musteri_listesi = {
                'df': df, 'tckn_sutun': tckn_sutun,
                'vkn_sutun': vkn_sutun, 'adsoyad_sutun': adsoyad_sutun
            }
            
            dosya_adi = os.path.basename(dosya_yolu)
            self.musteri_durum_label.config(text=f"✅ {dosya_adi} ({len(df)} kayıt)", fg='#2e7d32')
            self.toplu_musteri_label.config(text=f"✅ {dosya_adi} ({len(df)} kayıt)", fg='#2e7d32')
            self.durum_label.config(text=f"Müşteri listesi yüklendi: {len(df)} kayıt")
            
        except Exception as e:
            messagebox.showerror("Hata", f"Müşteri listesi yüklenemedi:\n{str(e)}")
    
    def musteri_sorgula(self, tckn=None, vkn=None):
        if self.musteri_listesi is None:
            return None, None
        
        df = self.musteri_listesi['df']
        tckn_s = self.musteri_listesi['tckn_sutun']
        vkn_s = self.musteri_listesi['vkn_sutun']
        ad_s = self.musteri_listesi['adsoyad_sutun']
        
        sonuc = None
        
        if tckn and tckn_s:
            eslesme = df[df[tckn_s].astype(str).str.strip() == str(tckn).strip()]
            if not eslesme.empty:
                sonuc = eslesme.iloc[0]
        
        if sonuc is None and vkn and vkn_s:
            eslesme = df[df[vkn_s].astype(str).str.strip() == str(vkn).strip()]
            if not eslesme.empty:
                sonuc = eslesme.iloc[0]
        
        if sonuc is not None:
            adsoyad = sonuc[ad_s] if ad_s and ad_s in sonuc.index else None
            return True, adsoyad
        
        return False, None

    # ==================== BELGE OLUŞTURMA ====================
    
    def belge_olustur(self, muhatap_kurum, muhatap_alt1, muhatap_alt2, tarih, sayi, tckn, vkn, adsoyad, musteri_durumu):
        """Word belgesi oluştur"""
        doc = Document()
        
        # Stil
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        # T.C.
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("T.C.").bold = True
        
        # Muhatap kurumları
        for kurum in [muhatap_kurum, muhatap_alt1, muhatap_alt2]:
            if kurum:
                pk = doc.add_paragraph()
                pk.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pk.add_run(kurum.upper()).bold = True
        
        # Tarih
        pt = doc.add_paragraph()
        pt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pt.add_run(tarih or datetime.now().strftime("%d.%m.%Y"))
        
        # İlgi
        doc.add_paragraph()
        doc.add_paragraph(f"İlgi: {tarih} tarihli ve {sayi} sayılı yazınız.")
        
        # Ana metin
        doc.add_paragraph()
        
        kimlik = tckn if tckn else vkn
        kimlik_tipi = "T.C. Kimlik Numaralı" if tckn else "Vergi Kimlik Numaralı"
        
        if "DEĞİL" in musteri_durumu.upper():
            metin = f"İlgi'de kayıtlı yazınıza istinaden Bankamız nezdinde gerekli araştırma yapılmış olup, {kimlik} {kimlik_tipi}"
            if adsoyad:
                metin += f" {adsoyad}'ın"
            metin += " Bankamız müşterisi olmadığı tespit edilmiştir."
        else:
            metin = f"İlgi'de kayıtlı yazınıza istinaden Bankamız nezdinde gerekli araştırma yapılmış olup, {kimlik} {kimlik_tipi}"
            if adsoyad:
                metin += f" {adsoyad}"
            metin += " ile ilgili gerekli işlemler yapılmaktadır."
        
        doc.add_paragraph(metin)
        
        # Kapanış
        doc.add_paragraph()
        doc.add_paragraph("Bilgilerinize arz ederiz.")
        doc.add_paragraph()
        doc.add_paragraph("Saygılarımızla,")
        doc.add_paragraph()
        
        pi1 = doc.add_paragraph()
        pi1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pi1.add_run("AYTEMİZ YATIRIM BANKASI A.Ş.").bold = True
        
        pi2 = doc.add_paragraph()
        pi2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pi2.add_run("Genel Müdürlük").bold = True
        
        return doc

    # ==================== TEKLİ İŞLEM ARAYÜZÜ ====================
    
    def tekli_arayuz_olustur(self):
        # Başlık
        baslik = tk.Frame(self.tekli_frame, bg='#1a237e', height=50)
        baslik.pack(fill=tk.X)
        baslik.pack_propagate(False)
        tk.Label(baslik, text="📄 Tekli Yazı İşleme", font=('Segoe UI', 14, 'bold'),
                bg='#1a237e', fg='white').pack(pady=12)
        
        # Müşteri listesi
        ml = tk.Frame(self.tekli_frame, bg='#e8f5e9', pady=5)
        ml.pack(fill=tk.X, padx=10, pady=5)
        tk.Button(ml, text="👥 Müşteri Listesi Yükle", command=self.musteri_listesi_sec,
                 font=('Segoe UI', 9), bg='#43A047', fg='white', cursor='hand2').pack(side=tk.LEFT, padx=5)
        self.musteri_durum_label = tk.Label(ml, text="❌ Müşteri listesi yüklenmedi", 
                                            font=('Segoe UI', 9), bg='#e8f5e9', fg='#c62828')
        self.musteri_durum_label.pack(side=tk.LEFT, padx=10)
        
        # Ana içerik
        ana = tk.Frame(self.tekli_frame, bg='#f0f0f0')
        ana.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # Sol panel
        sol = tk.LabelFrame(ana, text="📁 Gelen Yazı", font=('Segoe UI', 10, 'bold'), 
                           bg='#f0f0f0', padx=10, pady=10)
        sol.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0,5))
        
        # Dosya yükleme
        df = tk.Frame(sol, bg='#f0f0f0')
        df.pack(fill=tk.X)
        tk.Button(df, text="📄 Yazı Yükle (DOCX/PDF/TXT)", command=self.tekli_dosya_yukle,
                 font=('Segoe UI', 10), bg='#4CAF50', fg='white', cursor='hand2').pack(side=tk.LEFT)
        self.tekli_dosya_label = tk.Label(df, text="Dosya seçilmedi", 
                                          font=('Segoe UI', 9), bg='#f0f0f0', fg='#666')
        self.tekli_dosya_label.pack(side=tk.LEFT, padx=10)
        
        # Yazı içeriği
        tk.Label(sol, text="Yazı İçeriği:", font=('Segoe UI', 10, 'bold'), 
                bg='#f0f0f0').pack(anchor=tk.W, pady=(10,0))
        self.yazi_text = scrolledtext.ScrolledText(sol, height=10, font=('Consolas', 9))
        self.yazi_text.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # Bilgi alanları
        bilgi = tk.LabelFrame(sol, text="📝 Çıkarılan Bilgiler", font=('Segoe UI', 9, 'bold'),
                             bg='#f0f0f0', padx=10, pady=5)
        bilgi.pack(fill=tk.X, pady=5)
        
        self.entries = {}
        labels = [("Muhatap Kurum:", 0), ("Alt Birim 1:", 1), ("Alt Birim 2:", 2),
                  ("Yazı Tarihi:", 3), ("Sayı No:", 4), ("TCKN:", 5), ("VKN:", 6), ("Ad Soyad:", 7)]
        
        for lbl, row in labels:
            tk.Label(bilgi, text=lbl, bg='#f0f0f0').grid(row=row, column=0, sticky=tk.W, pady=1)
            e = tk.Entry(bilgi, width=50, font=('Segoe UI', 9))
            e.grid(row=row, column=1, pady=1, padx=5, sticky=tk.W)
            self.entries[lbl] = e
        
        # Müşteri sonuç
        self.musteri_sonuc = tk.Label(bilgi, text="", font=('Segoe UI', 9, 'bold'), bg='#f0f0f0')
        self.musteri_sonuc.grid(row=5, column=2, rowspan=2, padx=10)
        
        # Müşteri durumu
        tk.Label(bilgi, text="Müşteri Durumu:", bg='#f0f0f0').grid(row=8, column=0, sticky=tk.W, pady=1)
        self.musteri_combo = ttk.Combobox(bilgi, values=["Müşterimiz DEĞİL", "Müşterimiz - Manuel işlem"], 
                                          width=47, font=('Segoe UI', 9))
        self.musteri_combo.grid(row=8, column=1, pady=1, padx=5)
        self.musteri_combo.set("Müşterimiz DEĞİL")
        
        # Sağ panel - Önizleme
        sag = tk.LabelFrame(ana, text="📋 Yanıt Önizleme", font=('Segoe UI', 10, 'bold'),
                           bg='#f0f0f0', padx=10, pady=10)
        sag.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(5,0))
        
        self.onizleme_text = scrolledtext.ScrolledText(sag, height=20, font=('Consolas', 10), bg='#fffde7')
        self.onizleme_text.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # Butonlar
        btn = tk.Frame(sag, bg='#f0f0f0')
        btn.pack(fill=tk.X, pady=5)
        
        tk.Button(btn, text="🔍 Bilgileri Çıkar", command=self.tekli_bilgi_cikar,
                 font=('Segoe UI', 10, 'bold'), bg='#2196F3', fg='white', 
                 padx=15, cursor='hand2').pack(side=tk.LEFT, padx=3)
        tk.Button(btn, text="👁️ Önizle", command=self.tekli_onizle,
                 font=('Segoe UI', 10, 'bold'), bg='#9C27B0', fg='white',
                 padx=15, cursor='hand2').pack(side=tk.LEFT, padx=3)
        tk.Button(btn, text="💾 Kaydet", command=self.tekli_kaydet,
                 font=('Segoe UI', 10, 'bold'), bg='#4CAF50', fg='white',
                 padx=15, cursor='hand2').pack(side=tk.LEFT, padx=3)

    def tekli_dosya_yukle(self):
        dosya = filedialog.askopenfilename(
            title="TBB Yazısını Seçin",
            filetypes=[("Desteklenen Dosyalar", "*.docx *.pdf *.txt"), 
                      ("Word", "*.docx"), ("PDF", "*.pdf"), ("Metin", "*.txt")]
        )
        if dosya:
            self.tekli_dosya_label.config(text=os.path.basename(dosya))
            self.durum_label.config(text="Dosya okunuyor...")
            self.root.update()
            
            icerik = self.dosya_oku(dosya)
            self.yazi_text.delete('1.0', tk.END)
            self.yazi_text.insert('1.0', icerik)
            
            self.tekli_bilgi_cikar()
            self.durum_label.config(text=f"Yüklendi: {os.path.basename(dosya)}")
    
    def tekli_bilgi_cikar(self):
        icerik = self.yazi_text.get('1.0', tk.END)
        bilgiler = self.bilgi_cikar(icerik)
        
        alan_map = {
            'muhatap_kurum': "Muhatap Kurum:", 'muhatap_alt1': "Alt Birim 1:",
            'muhatap_alt2': "Alt Birim 2:", 'tarih': "Yazı Tarihi:",
            'sayi': "Sayı No:", 'tckn': "TCKN:", 'vkn': "VKN:", 'adsoyad': "Ad Soyad:"
        }
        
        for key, entry_key in alan_map.items():
            self.entries[entry_key].delete(0, tk.END)
            self.entries[entry_key].insert(0, bilgiler[key])
        
        # Müşteri kontrolü
        musteri_mi, adsoyad = self.musteri_sorgula(bilgiler['tckn'], bilgiler['vkn'])
        
        if self.musteri_listesi:
            if musteri_mi:
                self.musteri_sonuc.config(text="✅ MÜŞTERİMİZ", fg='#2e7d32')
                self.musteri_combo.set("Müşterimiz - Manuel işlem")
                if adsoyad and not bilgiler['adsoyad']:
                    self.entries["Ad Soyad:"].delete(0, tk.END)
                    self.entries["Ad Soyad:"].insert(0, str(adsoyad))
            else:
                self.musteri_sonuc.config(text="❌ MÜŞTERİMİZ DEĞİL", fg='#c62828')
                self.musteri_combo.set("Müşterimiz DEĞİL")
        else:
            self.musteri_sonuc.config(text="⚠️ Liste yüklenmedi", fg='#f57c00')
        
        self.tekli_onizle()
    
    def tekli_onizle(self):
        metin = "T.C.\n"
        
        for key in ["Muhatap Kurum:", "Alt Birim 1:", "Alt Birim 2:"]:
            val = self.entries[key].get().strip()
            if val:
                metin += f"{val.upper()}\n"
        
        tarih = self.entries["Yazı Tarihi:"].get() or datetime.now().strftime("%d.%m.%Y")
        sayi = self.entries["Sayı No:"].get()
        
        metin += f"\n{tarih}\n\n"
        metin += f"İlgi: {tarih} tarihli ve {sayi} sayılı yazınız.\n\n"
        
        tckn = self.entries["TCKN:"].get()
        vkn = self.entries["VKN:"].get()
        adsoyad = self.entries["Ad Soyad:"].get()
        
        kimlik = tckn if tckn else vkn
        kimlik_tipi = "T.C. Kimlik Numaralı" if tckn else "Vergi Kimlik Numaralı"
        
        if "DEĞİL" in self.musteri_combo.get().upper():
            metin += f"İlgi'de kayıtlı yazınıza istinaden Bankamız nezdinde gerekli araştırma yapılmış olup, {kimlik} {kimlik_tipi}"
            if adsoyad:
                metin += f" {adsoyad}'ın"
            metin += " Bankamız müşterisi olmadığı tespit edilmiştir.\n\n"
        else:
            metin += f"İlgi'de kayıtlı yazınıza istinaden Bankamız nezdinde gerekli araştırma yapılmış olup, {kimlik} {kimlik_tipi}"
            if adsoyad:
                metin += f" {adsoyad}"
            metin += " ile ilgili gerekli işlemler yapılmaktadır.\n\n"
        
        metin += "Bilgilerinize arz ederiz.\n\nSaygılarımızla,\n\n"
        metin += "AYTEMİZ YATIRIM BANKASI A.Ş.\nGenel Müdürlük"
        
        self.onizleme_text.delete('1.0', tk.END)
        self.onizleme_text.insert('1.0', metin)
    
    def tekli_kaydet(self):
        tckn = self.entries["TCKN:"].get().strip()
        vkn = self.entries["VKN:"].get().strip()
        
        if not tckn and not vkn:
            messagebox.showwarning("Uyarı", "TCKN veya VKN bilgisi gerekli!")
            return
        
        belge = self.belge_olustur(
            self.entries["Muhatap Kurum:"].get(),
            self.entries["Alt Birim 1:"].get(),
            self.entries["Alt Birim 2:"].get(),
            self.entries["Yazı Tarihi:"].get(),
            self.entries["Sayı No:"].get(),
            tckn, vkn,
            self.entries["Ad Soyad:"].get(),
            self.musteri_combo.get()
        )
        
        kimlik = tckn if tckn else vkn
        dosya = filedialog.asksaveasfilename(
            title="Yanıt Yazısını Kaydet",
            defaultextension=".docx",
            initialfile=f"Yanit_{kimlik}_{datetime.now().strftime('%Y%m%d')}.docx",
            filetypes=[("Word Belgesi", "*.docx")]
        )
        
        if dosya:
            try:
                belge.save(dosya)
                messagebox.showinfo("Başarılı", f"Yanıt kaydedildi:\n{dosya}")
                self.durum_label.config(text=f"Kaydedildi: {os.path.basename(dosya)}")
            except Exception as e:
                messagebox.showerror("Hata", f"Kayıt başarısız:\n{str(e)}")

    # ==================== TOPLU İŞLEM ARAYÜZÜ ====================
    
    def toplu_arayuz_olustur(self):
        # Başlık
        baslik = tk.Frame(self.toplu_frame, bg='#1565c0', height=50)
        baslik.pack(fill=tk.X)
        baslik.pack_propagate(False)
        tk.Label(baslik, text="📁 Toplu Yazı İşleme - Klasör Seç & Otomatik İşle",
                font=('Segoe UI', 14, 'bold'), bg='#1565c0', fg='white').pack(pady=12)
        
        # Kontroller
        ctrl = tk.Frame(self.toplu_frame, bg='#e3f2fd', pady=8)
        ctrl.pack(fill=tk.X, padx=10, pady=5)
        
        tk.Button(ctrl, text="👥 Müşteri Listesi", command=self.musteri_listesi_sec,
                 font=('Segoe UI', 9), bg='#43A047', fg='white', cursor='hand2').pack(side=tk.LEFT, padx=5)
        self.toplu_musteri_label = tk.Label(ctrl, text="❌ Yüklenmedi", 
                                            font=('Segoe UI', 9), bg='#e3f2fd', fg='#c62828')
        self.toplu_musteri_label.pack(side=tk.LEFT, padx=10)
        
        tk.Button(ctrl, text="📂 Klasör Seç", command=self.klasor_sec,
                 font=('Segoe UI', 11, 'bold'), bg='#1976D2', fg='white', 
                 padx=20, cursor='hand2').pack(side=tk.RIGHT, padx=5)
        self.klasor_label = tk.Label(ctrl, text="Klasör seçilmedi", 
                                     font=('Segoe UI', 9), bg='#e3f2fd', fg='#666')
        self.klasor_label.pack(side=tk.RIGHT, padx=10)
        
        # Treeview
        tree_frame = tk.Frame(self.toplu_frame, bg='#f0f0f0')
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        cols = ('dosya', 'tckn', 'vkn', 'adsoyad', 'sayi', 'musteri', 'durum')
        self.tree = ttk.Treeview(tree_frame, columns=cols, show='headings', height=18)
        
        basliklar = [('dosya', '📄 Dosya', 200), ('tckn', 'TCKN', 100), ('vkn', 'VKN', 90),
                     ('adsoyad', 'Ad Soyad', 130), ('sayi', 'Sayı No', 180), 
                     ('musteri', 'Müşteri?', 80), ('durum', 'Durum', 120)]
        
        for col, text, width in basliklar:
            self.tree.heading(col, text=text)
            self.tree.column(col, width=width)
        
        sb = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscroll=sb.set)
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Butonlar
        btn = tk.Frame(self.toplu_frame, bg='#f0f0f0', pady=10)
        btn.pack(fill=tk.X, padx=10)
        
        butonlar = [
            ("🔍 Tümünü Analiz Et", self.toplu_analiz, '#2196F3'),
            ("📝 Yanıtları Oluştur", self.toplu_yanit_olustur, '#4CAF50'),
            ("💾 Tümünü Kaydet", self.toplu_kaydet, '#FF9800'),
            ("📊 Rapor Çıkar", self.rapor_olustur, '#9C27B0')
        ]
        
        for text, cmd, color in butonlar:
            tk.Button(btn, text=text, command=cmd, font=('Segoe UI', 11, 'bold'),
                     bg=color, fg='white', padx=20, cursor='hand2').pack(side=tk.LEFT, padx=5)
        
        self.istatistik_label = tk.Label(btn, text="", font=('Segoe UI', 10), bg='#f0f0f0')
        self.istatistik_label.pack(side=tk.RIGHT, padx=20)
    
    def klasor_sec(self):
        klasor = filedialog.askdirectory(title="TBB Yazılarının Bulunduğu Klasörü Seçin")
        if not klasor:
            return
        
        self.secili_klasor = klasor
        self.klasor_label.config(text=os.path.basename(klasor))
        
        self.toplu_dosyalar = []
        for dosya in os.listdir(klasor):
            if os.path.splitext(dosya)[1].lower() in ['.docx', '.pdf', '.txt']:
                self.toplu_dosyalar.append(os.path.join(klasor, dosya))
        
        # Tabloyu temizle ve doldur
        for item in self.tree.get_children():
            self.tree.delete(item)
        
        for dosya in self.toplu_dosyalar:
            self.tree.insert('', tk.END, values=(os.path.basename(dosya), '', '', '', '', '', 'Bekliyor'))
        
        self.istatistik_label.config(text=f"📁 {len(self.toplu_dosyalar)} dosya bulundu")
        self.durum_label.config(text=f"Klasör seçildi: {len(self.toplu_dosyalar)} dosya")
    
    def toplu_analiz(self):
        if not self.toplu_dosyalar:
            messagebox.showwarning("Uyarı", "Önce bir klasör seçin!")
            return
        
        self.toplu_sonuclar = []
        self.durum_label.config(text="Analiz yapılıyor...")
        
        for i, dosya_yolu in enumerate(self.toplu_dosyalar):
            try:
                self.root.update()
                icerik = self.dosya_oku(dosya_yolu)
                bilgiler = self.bilgi_cikar(icerik)
                
                musteri_mi, adsoyad_db = self.musteri_sorgula(bilgiler['tckn'], bilgiler['vkn'])
                
                if self.musteri_listesi is None:
                    musteri_str = "⚠️"
                elif musteri_mi:
                    musteri_str = "✅ Evet"
                else:
                    musteri_str = "❌ Hayır"
                
                if adsoyad_db and not bilgiler['adsoyad']:
                    bilgiler['adsoyad'] = str(adsoyad_db)
                
                sonuc = {
                    'dosya': dosya_yolu,
                    'dosya_adi': os.path.basename(dosya_yolu),
                    'bilgiler': bilgiler,
                    'musteri_mi': musteri_mi
                }
                self.toplu_sonuclar.append(sonuc)
                
                item = self.tree.get_children()[i]
                self.tree.item(item, values=(
                    os.path.basename(dosya_yolu),
                    bilgiler['tckn'],
                    bilgiler['vkn'],
                    bilgiler['adsoyad'][:20] if bilgiler['adsoyad'] else '',
                    bilgiler['sayi'][:25] if bilgiler['sayi'] else '',
                    musteri_str,
                    "✓ Tamam"
                ))
                
            except Exception as e:
                item = self.tree.get_children()[i]
                self.tree.item(item, values=(os.path.basename(dosya_yolu), '', '', '', '', '', f"❌ Hata"))
        
        # İstatistik
        musteri_sayisi = sum(1 for s in self.toplu_sonuclar if s['musteri_mi'])
        degil_sayisi = len(self.toplu_sonuclar) - musteri_sayisi
        
        self.istatistik_label.config(
            text=f"📊 Toplam: {len(self.toplu_sonuclar)} | ✅ Müşteri: {musteri_sayisi} | ❌ Değil: {degil_sayisi}"
        )
        self.durum_label.config(text="Analiz tamamlandı")
        messagebox.showinfo("Tamamlandı", f"{len(self.toplu_sonuclar)} dosya analiz edildi!\n\n"
                           f"✅ Müşteri: {musteri_sayisi}\n❌ Müşteri değil: {degil_sayisi}")
    
    def toplu_yanit_olustur(self):
        if not self.toplu_sonuclar:
            messagebox.showwarning("Uyarı", "Önce analiz yapın!")
            return
        
        self.toplu_yanitlar = []
        
        for sonuc in self.toplu_sonuclar:
            if not sonuc['musteri_mi']:
                b = sonuc['bilgiler']
                belge = self.belge_olustur(
                    b['muhatap_kurum'], b['muhatap_alt1'], b['muhatap_alt2'],
                    b['tarih'], b['sayi'], b['tckn'], b['vkn'], b['adsoyad'],
                    "Müşterimiz DEĞİL"
                )
                self.toplu_yanitlar.append({
                    'dosya_adi': sonuc['dosya_adi'],
                    'tckn': b['tckn'],
                    'vkn': b['vkn'],
                    'belge': belge
                })
        
        self.durum_label.config(text=f"{len(self.toplu_yanitlar)} yanıt oluşturuldu")
        messagebox.showinfo("Tamamlandı", f"{len(self.toplu_yanitlar)} yanıt yazısı oluşturuldu!\n\n"
                           "(Sadece müşteri olmayanlar için)")
    
    def toplu_kaydet(self):
        if not self.toplu_yanitlar:
            messagebox.showwarning("Uyarı", "Önce yanıtları oluşturun!")
            return
        
        klasor = filedialog.askdirectory(title="Yanıtların Kaydedileceği Klasörü Seçin")
        if not klasor:
            return
        
        basarili = 0
        for yanit in self.toplu_yanitlar:
            try:
                kimlik = yanit['tckn'] if yanit['tckn'] else yanit['vkn']
                dosya_adi = f"Yanit_{kimlik}_{datetime.now().strftime('%Y%m%d')}.docx"
                dosya_yolu = os.path.join(klasor, dosya_adi)
                yanit['belge'].save(dosya_yolu)
                basarili += 1
            except:
                pass
        
        self.durum_label.config(text=f"{basarili} yanıt kaydedildi")
        messagebox.showinfo("Tamamlandı", f"{basarili} yanıt yazısı kaydedildi!\n\nKonum: {klasor}")
    
    def rapor_olustur(self):
        if not self.toplu_sonuclar:
            messagebox.showwarning("Uyarı", "Önce analiz yapın!")
            return
        
        dosya = filedialog.asksaveasfilename(
            title="Raporu Kaydet",
            defaultextension=".xlsx",
            initialfile=f"TBB_Rapor_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
            filetypes=[("Excel", "*.xlsx")]
        )
        
        if not dosya:
            return
        
        try:
            data = []
            for s in self.toplu_sonuclar:
                b = s['bilgiler']
                data.append({
                    'Dosya': s['dosya_adi'],
                    'Muhatap Kurum': b['muhatap_kurum'],
                    'Tarih': b['tarih'],
                    'Sayı': b['sayi'],
                    'TCKN': b['tckn'],
                    'VKN': b['vkn'],
                    'Ad Soyad': b['adsoyad'],
                    'Müşteri mi?': 'Evet' if s['musteri_mi'] else 'Hayır',
                    'Aksiyon': 'Manuel işlem gerekli' if s['musteri_mi'] else 'Otomatik yanıt oluşturuldu'
                })
            
            pd.DataFrame(data).to_excel(dosya, index=False)
            messagebox.showinfo("Tamamlandı", f"Rapor kaydedildi:\n{dosya}")
            self.durum_label.config(text=f"Rapor: {os.path.basename(dosya)}")
        except Exception as e:
            messagebox.showerror("Hata", f"Rapor oluşturulamadı:\n{str(e)}")


def main():
    root = tk.Tk()
    app = TBBYanitSistemi(root)
    root.mainloop()


if __name__ == "__main__":
    main()
